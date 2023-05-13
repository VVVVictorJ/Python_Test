import os
import re
import uuid

from docx import Document


class ProcessDocx:
    docxWordInfo = {}

    def writeToFile(self, source, dest, filename: str):
        """写入文件至指定位置

        Args:
            source (binary): 源文件
            dest (str): 目标地址
            filename (str): 文件名
        """
        destinationPath = os.path.sep.join([dest, filename])
        try:
            with open(f"{destinationPath}", "wb") as f:
                f.write(source)
        except FileNotFoundError:
            print("无法打开指定的文件!")
        except LookupError:
            print("指定了未知的编码!")
        except UnicodeDecodeError:
            print("读取文件时解码错误!")

    def saveToDict(self, k, v):
        """存储数据至字典

        Args:
            k (str): Key
            v (str): value
        """
        self.docxWordInfo[k] = [v]

    def removeDuplication(self, cellList: list):
        """模拟入栈去除重复数值

        Args:
            cellList (list): docx表格所有信息

        Returns:
            _type_: list
        """
        cells_value = []
        for ele in cellList:
            # 栈为空直接入栈
            if len(cells_value) == 0:
                cells_value.append(ele)
            else:
                if cells_value[-1] != ele:
                    cells_value.append(ele)
        return cells_value

    def processDocxTableInfo(self, filepath: str):
        """处理Docx表格数据

        Args:
            filepath (str): docx所在路径

        Returns:
            _type_: dict
        """
        try:
            file = Document(filepath)
            table = file.tables[0]
            cells = table._cells
            cells_string = [cell.text for cell in cells]
            cells_value = self.removeDuplication(cells_string)
            docxInfo = {}

            def isR(l: list):
                for i in l:
                    if "√" in i:
                        return i

            # 故障系统
            self.saveToDict(cells_value[0], cells_value[3])
            # 故障名称
            self.saveToDict(cells_value[1], cells_value[4])
            # 签收日期
            self.saveToDict(cells_value[2], cells_value[5])
            # 故障日期
            self.saveToDict(cells_value[6], cells_value[10])
            # 处理起止日期和时间 去除换行
            self.saveToDict(cells_value[7], cells_value[11].replace("\n", " "))
            # 填表人
            self.saveToDict(cells_value[8], cells_value[12])
            # 负责人
            self.saveToDict(cells_value[9], cells_value[13])
            # 主机系统停机时间
            self.saveToDict(cells_value[14], cells_value[18])
            # [19] 存在年月
            if "月" in cells_value[19]:
                # 历时
                self.saveToDict("历时一", cells_value[18])
                # 影响业务时间
                self.saveToDict(cells_value[16], cells_value[19].replace("\n", " "))
                # 历时
                self.saveToDict("历时二", cells_value[20])
                # 现象、排障经过、原因分析 故障说明
                self.saveToDict(
                    cells_value[21].replace("\n", "、").replace(" ", ""),
                    cells_value[22].replace("\n", " "),
                )
                # 消耗备件
                self.saveToDict(cells_value[23][0:4], cells_value[23][-1])
                # 故障处理人
                self.saveToDict(cells_value[24][0:5], cells_value[24][6:])
                x = lambda x: isR(cells_value)
                self.saveToDict(
                    cells_value[cells_value.index("责任")], x(cells_value)[-4:]
                )
                self.saveToDict(
                    cells_value[cells_value.index("纠正和预防措施")],
                    cells_value[cells_value.index("纠正和预防措施") + 1].replace("\n", " "),
                )
            else:
                # 历时
                self.saveToDict("历时一", cells_value[19])
                # 影响业务时间
                self.saveToDict(cells_value[16], cells_value[20].replace("\n", " "))
                # 历时
                self.saveToDict("历时二", cells_value[21])
                # 现象、排障经过、原因分析 故障说明
                self.saveToDict(
                    cells_value[22].replace("\n", "、").replace(" ", ""),
                    cells_value[23].replace("\n", " "),
                )
                # 消耗备件
                self.saveToDict(cells_value[24][0:4], cells_value[24][-1])
                # 故障处理人
                self.saveToDict(cells_value[25][0:5], cells_value[25][6:])
                x = lambda x: isR(cells_value)
                self.saveToDict(
                    cells_value[cells_value.index("责任")], x(cells_value)[-4:]
                )
                self.saveToDict(
                    cells_value[cells_value.index("纠正和预防措施")],
                    cells_value[cells_value.index("纠正和预防措施") + 1].replace("\n", " "),
                )
        except FileNotFoundError:
            print("无法打开指定的文件!")
        finally:
            return self.docxWordInfo

    @staticmethod
    def processDocxImageInfo(filepath: str):
        ImageDict = {}
        try:
            doc = Document(filepath)
            dict_rel = doc.part._rels
            for rel in dict_rel:
                rel = dict_rel[rel]
                if "image" in rel.target_ref:
                    img_name = re.findall("/(.*)", rel.target_ref)[0]
                    word_name = os.path.splitext(filepath)[0]
                    if os.sep in word_name:
                        new_name = word_name.split("\\")[-1]
                    else:
                        new_name = word_name.split("/")[-1]
                    img_name = "{}-{}{}".format(
                        f"{new_name}", str(uuid.uuid4()), os.path.splitext(img_name)[-1]
                    )
                    ImageDict[img_name] = rel.target_part.blob
            # print(len(ImageDict))
            return ImageDict
        except:
            print("写图片出错")
        finally:
            return ImageDict

if __name__ == "__main__":
    obj = ProcessDocx()
    print(
        ProcessDocx.processDocxImageInfo(
            filepath=f"E:\\code\\Python\\shitwork\\src\\2518_20220110_湛江钢铁_炼钢热轧L3系统_故障报告书_B.docx",
        )
    )