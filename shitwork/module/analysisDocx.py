import os
import re
import uuid
import zipfile

from docx import Document
from lxml import etree
from recordFileList import RecordDirList


#TODO 将其实现成单例
class AnalysisDocx:

    def __init__(self, filepath: str):
        if not os.path.isfile(filepath):
            raise TypeError("NOT FILEPATH")
        if not filepath.endswith('.docx'):
            raise TypeError("NOT DOCX FILE")
        self.filepath = filepath
        self.uuidStr = str(uuid.uuid4())
        self.resultpath = ""
        self.docxWordInfo = {}
        self.docxImage = []
        self.record = RecordDirList()

    def genData(self, k, v):
        self.docxWordInfo[k] = v

    def genImageRocord(self, val):
        self.docxImage.append(val)

    def genDocxTableInfo2TxtTest(self):
        try:
            file = Document(self.filepath)
            table = file.tables[0]
            cells = table._cells
            cells_string = [cell.text for cell in cells]
            # print(len(cells_string))
            # for k, v in enumerate(cells_string):
            #     print(k, v.replace('\n', ' ').replace(' ', ''))
            cells_value = []
            cnt = 0
            for ele in cells_string:
                if (len(cells_value) == 0):
                    cells_value.append(ele)
                else:
                    # if ele == '无':
                    #     cnt += 1
                    if (cells_value[-1] != ele):
                        cells_value.append(ele)
                    # elif cells_value[-1] == '无' and cnt == 7:
                    #     cells_value.append(ele)
            # print(len(cells_value))
            # cells_value = sorted(set(cells_string), key=cells_string.index)
            # print(len(cells_value))
            for k, v in enumerate(cells_value):
                print(k, v.replace('\n', ' '))
        except:
            pass

    def genDocxTableInfo2Txt(self):
        try:
            file = Document(self.filepath)
            table = file.tables[0]
            cells = table._cells
            cells_string = [cell.text for cell in cells]
            # cells_value = sorted(set(cells_string), key=cells_string.index)
            cells_value = []
            for ele in cells_string:
                if (len(cells_value) == 0):
                    cells_value.append(ele)
                else:
                    if (cells_value[-1] != ele):
                        cells_value.append(ele)
            def isR(l: list):
                for i in l:
                    if '√' in i:
                        return i
            # 故障系统
            self.genData(cells_value[0], cells_value[3])
            # 故障名称
            self.genData(cells_value[1], cells_value[4])
            # 签收日期
            self.genData(cells_value[2], cells_value[5])
            # 故障日期
            self.genData(cells_value[6], cells_value[10])
            # 处理起止日期和时间 去除换行
            self.genData(cells_value[7], cells_value[11].replace('\n', ' '))
            # 填表人
            self.genData(cells_value[8], cells_value[12])
            # 负责人
            self.genData(cells_value[9], cells_value[13])
            # 主机系统停机时间
            self.genData(cells_value[14], cells_value[18])
            # [19] 存在年月 
            if('月' in cells_value[19]):
                # 历时
                self.genData('历时一', cells_value[18])
                # 影响业务时间
                self.genData(cells_value[16], cells_value[19].replace('\n', ' '))
                # 历时
                self.genData('历时二', cells_value[20])
                #现象、排障经过、原因分析 故障说明
                self.genData(cells_value[21].replace('\n', '、').replace(' ',''), cells_value[22].replace('\n', ' '))
                # 消耗备件
                self.genData(cells_value[23][0:4], cells_value[23][-1])   
                # 故障处理人
                self.genData(cells_value[24][0:5], cells_value[24][6:])     
                x = lambda x: isR(cells_value)     
                self.genData(cells_value[cells_value.index('责任')],
                x(cells_value)[-4:])   
                self.genData(
                cells_value[cells_value.index('纠正和预防措施')],
                cells_value[cells_value.index('纠正和预防措施') + 1].replace(
                    '\n', ' '))
            else:
                # 历时
                self.genData('历时一', cells_value[19])
                # 影响业务时间
                self.genData(cells_value[16], cells_value[20].replace('\n', ' '))
                # 历时
                self.genData('历时二', cells_value[21])
                #现象、排障经过、原因分析 故障说明
                self.genData(cells_value[22].replace('\n', '、').replace(' ',''), cells_value[23].replace('\n', ' '))
                # 消耗备件
                self.genData(cells_value[24][0:4], cells_value[24][-1])   
                # 故障处理人
                self.genData(cells_value[25][0:5], cells_value[25][6:])     
                x = lambda x: isR(cells_value)     
                self.genData(cells_value[cells_value.index('责任')],
                x(cells_value)[-4:])   
                self.genData(
                cells_value[cells_value.index('纠正和预防措施')],
                cells_value[cells_value.index('纠正和预防措施') + 1].replace(
                    '\n', ' '))
        except:
            print('出错')
        try:
            txtFileName = "{}/{}{}".format(self.resultpath, self.uuidStr,
                                           '.txt')
            with open(txtFileName, "wb") as f:
                for k, v in self.docxWordInfo.items():
                    f.write("{0:{1}<10}\t{1:^10}@{2}\t{3:^10}\n".format(
                        k, chr(12288), v, chr(12288)).encode())
        except FileNotFoundError:
            print('无法打开指定的文件!')
        except LookupError:
            print('指定了未知的编码!')
        except UnicodeDecodeError:
            print('读取文件时解码错误!')

    def genDocxImage(self):
        self.uuidStr = str(uuid.uuid4())
        self.resultpath = os.path.join(os.getcwd(), self.uuidStr)
        try:
            doc = Document(self.filepath)
            dict_rel = doc.part._rels
            for rel in dict_rel:
                rel = dict_rel[rel]
                if "image" in rel.target_ref:
                    if not os.path.exists(self.resultpath):
                        os.makedirs(self.resultpath)
                    img_name = re.findall("/(.*)", rel.target_ref)[0]
                    word_name = os.path.splitext(self.filepath)[0]
                    if os.sep in word_name:
                        new_name = word_name.split('\\')[-1]
                    else:
                        new_name = word_name.split('/')[-1]
                    img_name = "{}-{}{}".format(f'{new_name}',
                                                str(uuid.uuid4()),
                                                os.path.splitext(img_name)[-1])
                    with open(f'{self.resultpath}/{img_name}', "wb") as f:
                        f.write(rel.target_part.blob)
                        self.genImageRocord(f'{self.resultpath}/{img_name}')
        except:
            print('写图片出错')

    def showData(self):
        for k, v in self.docxWordInfo.items():
            print("{0:{1}<10}\t{1:^10}".format(k, chr(12288)),
                  "{0}\t{1:^10}".format(v, chr(12288)))
        for k, v in enumerate(self.docxImage):
            print("{0:{1}<2}\t{1:^3}".format(k, chr(12288)),
                  "{0}\t{1:^10}".format(v, chr(12288)))


if __name__ == '__main__':
    obj = AnalysisDocx(filepath=f"a.docx")
    # 有待封装，将其封装成一个类Producer
    # obj.genDocxTableInfo2TxtTest()
    # obj.genDocxTableInfo2Txt()
    # obj.showData()
    try:
        # obj.genDocxTableInfo2TxtTest()
        obj.genDocxImage()
        obj.genDocxTableInfo2Txt()
        # obj.showData()
    except:
        pass
    finally:
        obj.record.addList(obj.resultpath)
